""" Document Service- extend text drom files"""
import logging

from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger(__name__)

class DocumentParser:
    @staticmethod
    def parse_pdf(file_path):
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f'--- Page {page_num} ---\n{page_text}')
            
            full_text = '\n\n'.join(text_parts)
            word_count = len(full_text.split())
            
            return {
                'text': full_text,
                'page_count': len(reader.pages),
                'word_count': word_count,
            }
        except Exception as e :
            logger.error(f'PDF parsing error: {e}')
            raise ValueError(f' Do not read PDF file: {str(e)}')
        
    @staticmethod
    def parse_docx(file_path):
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            #paragrphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            #schedules
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)

                full_text = '\n'.join(text_parts)
                word_count = len(full_text.split())

                return {
                    'text': full_text,
                    'page_count': 1,  # DOCX'da sahifa aniq emas
                    'word_count': word_count,
                }
            
        except Exception as e:
            logger.error(f'DOCX parsing error: {e}')
            raise ValueError(f' Do not read Word file: {str(e)}')
        
    @staticmethod
    def parse_txt(file_path):
        try:
            #first try with UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='cp1251') as f:
                    text = f.read()
            
            word_count = len(text.split())

            return{
                'text': text,
                'page_count': 1,
                'word_count': word_count,
            }
        except Exception as e:
            logger.error(f'TXT parsing error: {e}')
            raise ValueError(f' Do not read Text file: {str(e)}')

class DocumentProcessor:
    @staticmethod
    def process_document(document):
        from .models import Document

        try:
            document.status = Document.Status.PROCESSING
            document.save(update_fields=['status'])

            file_path = document.file.path
            file_type = document.file_type

            if file_type ==Document.FileType.PDF:
                result = DocumentParser.parse_pdf(file_path)
            elif file_type == Document.FileType.DOCX:
                result = DocumentParser.parse_docx(file_path)
            elif file_type == Document.FileType.TXT:
                result = DocumentParser.parse_txt(file_path)
            else:
                raise ValueError(f' Invalid file type: {file_type}')
            
            #save result
            document.extracted_text = result['text']
            document.page_count = result['page_count']
            document.word_count = result['word_count']
            document.status = Document.Status.COMPLETED
            document.error_message =''
            document.save()

            logger.info(f'Document{document.id} successfully proccessed')
            return document
        
        except Exception as e:
            document.status = Document.Status.FAILED
            document.error_message = str(e)
            document.save(update_fields=['status', 'error_message'])

            logger.erro(f'Document {document.id} proccessing failed: {e}')
            raise


            





